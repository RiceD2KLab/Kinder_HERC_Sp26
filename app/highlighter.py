"""
app/highlighter.py
Builds a highlighted .docx transcript from prediction results.
predicted_label == 1 → yellow highlight
predicted_label == 0 → plain text
"""

from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def build_docx(
    predictions: List[Dict],
    output_path: Path,
    district: str = "",
    video_url: str = "",
) -> Path:
    doc = Document()

    # Page: US Letter, 1-inch margins
    sec = doc.sections[0]
    sec.page_width    = int(8.5 * 914400)
    sec.page_height   = int(11  * 914400)
    sec.left_margin   = sec.right_margin = sec.top_margin = sec.bottom_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    r = title.add_run("Highlighted Transcript")
    r.bold           = True
    r.font.size      = Pt(18)
    r.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

    # District / source
    meta = doc.add_paragraph()
    if district:
        _bold_line(meta, "District: ", district)
    if video_url:
        _bold_line(meta, "Source: ", video_url)

    doc.add_paragraph()

    # Chunks
    for chunk in predictions:
        flagged = chunk.get("predicted_label") == 1

        # Timestamp
        ts = doc.add_paragraph()
        ts.paragraph_format.space_before = Pt(8)
        ts.paragraph_format.space_after  = Pt(2)
        ts_r = ts.add_run(f"[{chunk.get('window_start','')} – {chunk.get('window_end','')}]")
        ts_r.font.size      = Pt(9)
        ts_r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Body
        body = doc.add_paragraph()
        body.paragraph_format.space_after = Pt(4)
        body_r = body.add_run(chunk.get("text", "").strip())
        body_r.font.size = Pt(11)

        if flagged:
            body_r.bold = True
            _shade(body_r, "FFFF99")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def _bold_line(para, label, value):
    lb = para.add_run(label)
    lb.bold = True
    lb.font.size = Pt(10)
    vr = para.add_run(value + "\n")
    vr.font.size = Pt(10)


def _shade(run, hex_fill):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_fill)
    rPr.append(shd)